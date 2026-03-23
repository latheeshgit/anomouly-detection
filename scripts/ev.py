import os
import pandas as pd
import torch
import torch.nn as nn
from sklearn.preprocessing import MinMaxScaler
from sklearn.metrics import (
    confusion_matrix, precision_score, recall_score, f1_score, roc_auc_score
)
from sklearn.model_selection import train_test_split
from docx import Document
import numpy as np

# Paths
model_dir = r"E:\CloudAnomalyDetectionSystem\models"
test_dir = r"E:\CloudAnomalyDetectionSystem\data\test"
results_dir = r"E:\CloudAnomalyDetectionSystem\results\evaluation_metrics"
os.makedirs(results_dir, exist_ok=True)

# Features
features = [
    'Destination Port', 'Flow Duration', 'Total Fwd Packets', 'Total Backward Packets',
    'Total Length of Fwd Packets', 'Total Length of Bwd Packets', 'Fwd Packet Length Max',
    'Fwd Packet Length Min', 'Fwd Packet Length Mean', 'Fwd Packet Length Std', 'Bwd Packet Length Max',
    'Bwd Packet Length Min', 'Bwd Packet Length Mean', 'Bwd Packet Length Std', 'Flow Bytes/s',
    'Flow Packets/s', 'Flow IAT Mean', 'Flow IAT Std', 'Flow IAT Max', 'Flow IAT Min', 'Fwd IAT Total',
    'Fwd IAT Mean', 'Fwd IAT Std', 'Fwd IAT Max', 'Fwd IAT Min', 'Bwd IAT Total', 'Bwd IAT Mean',
    'Bwd IAT Std', 'Bwd IAT Max', 'Bwd IAT Min', 'Fwd PSH Flags', 'Bwd PSH Flags', 'Fwd URG Flags',
    'Bwd URG Flags', 'Fwd Header Length', 'Bwd Header Length', 'Fwd Packets/s', 'Bwd Packets/s',
    'Min Packet Length', 'Max Packet Length', 'Packet Length Mean', 'Packet Length Std',
    'Packet Length Variance', 'FIN Flag Count', 'SYN Flag Count', 'RST Flag Count', 'PSH Flag Count',
    'ACK Flag Count', 'URG Flag Count', 'CWE Flag Count', 'ECE Flag Count', 'Down/Up Ratio',
    'Average Packet Size', 'Avg Fwd Segment Size', 'Avg Bwd Segment Size',
    'Fwd Header Length', 'Fwd Avg Bytes/Bulk', 'Fwd Avg Packets/Bulk', 'Fwd Avg Bulk Rate',
    'Bwd Avg Bytes/Bulk', 'Bwd Avg Packets/Bulk', 'Bwd Avg Bulk Rate', 'Subflow Fwd Packets',
    'Subflow Fwd Bytes', 'Subflow Bwd Packets', 'Subflow Bwd Bytes', 'Init_Win_bytes_forward',
    'Init_Win_bytes_backward', 'act_data_pkt_fwd', 'min_seg_size_forward', 'Active Mean',
    'Active Std', 'Active Max', 'Active Min', 'Idle Mean', 'Idle Std', 'Idle Max', 'Idle Min',
    'Label'
]

# Load test data
def load_data(directory):
    dataframes = []
    for file in os.listdir(directory):
        if file.endswith(".csv"):
            file_path = os.path.join(directory, file)
            df = pd.read_csv(file_path, usecols=features)
            dataframes.append(df)
    return pd.concat(dataframes, ignore_index=True)

test_data = load_data(test_dir)

# Handle missing labels
print(f"Before dropping NaN labels: {test_data.shape}")
label_mapping = {
    'BENIGN': 0, 'DDoS': 1, 'PortScan': 1, 'Bot': 1, 'Infiltration': 1,
    'Web Attack - XSS': 1, 'Web Attack - Brute Force': 1, 'Web Attack - Sql Injection': 1,
    'FTP-Patator': 1, 'SSH-Patator': 1, 'DoS Hulk': 1, 'DoS GoldenEye': 1,
    'DoS slowloris': 1, 'DoS Slowhttptest': 1, 'Heartbleed': 1
}
test_data['Label'] = test_data['Label'].map(label_mapping)
if test_data['Label'].isnull().any():
    print("Warning: NaN values detected in labels. Rows with NaN labels will be dropped.")
    test_data = test_data.dropna(subset=['Label'])
print(f"After dropping NaN labels: {test_data.shape}")

# Check for missing values in features
print(test_data.isnull().sum())  # Show count of missing values per column
if test_data.isnull().any().any():
    test_data = test_data.dropna()  # Drop rows with any NaN values in features
    print(f"After dropping rows with NaN values in features: {test_data.shape}")

# Ensure sufficient data remains
if test_data.shape[0] == 0:
    raise ValueError("No data left after handling missing values!")

# Preprocess data
scaler = MinMaxScaler()
X = scaler.fit_transform(test_data.drop(columns=['Label']))
y = test_data['Label'].values
X_train, X_val, y_train, y_val = train_test_split(X, y, test_size=0.2, random_state=42)

# Convert to PyTorch tensors
X_train_tensor = torch.tensor(X_train, dtype=torch.float32)
y_train_tensor = torch.tensor(y_train, dtype=torch.long)
X_val_tensor = torch.tensor(X_val, dtype=torch.float32)
y_val_tensor = torch.tensor(y_val, dtype=torch.long)

# Autoencoder Model
class Autoencoder(nn.Module):
    def __init__(self, input_dim):
        super(Autoencoder, self).__init__()
        self.encoder = nn.Sequential(
            nn.Linear(input_dim, 64),
            nn.ReLU(),
            nn.Linear(64, 32),
            nn.ReLU()
        )
        self.decoder = nn.Sequential(
            nn.Linear(32, 64),
            nn.ReLU(),
            nn.Linear(64, input_dim),
            nn.Sigmoid()
        )

    def forward(self, x):
        encoded = self.encoder(x)
        decoded = self.decoder(encoded)
        return decoded

# Load Autoencoder Model
autoencoder = Autoencoder(input_dim=X_train.shape[1])
autoencoder_path = r"E:\CloudAnomalyDetectionSystem\models\autoencoder_model.pth"
if not os.path.exists(autoencoder_path):
    raise FileNotFoundError(f"Autoencoder model file not found at {autoencoder_path}")
autoencoder.load_state_dict(torch.load(autoencoder_path, weights_only=True))

# CNN Model
class CNNModel(nn.Module):
    def __init__(self, input_dim, num_classes=1):  # Single output for binary classification
        super(CNNModel, self).__init__()
        self.conv1 = nn.Conv1d(1, 16, kernel_size=3, stride=1, padding=1)
        self.conv2 = nn.Conv1d(16, 32, kernel_size=3, stride=1, padding=1)
        self.flatten = nn.Flatten()
        self.fc1 = nn.Linear(32 * input_dim, 128)
        self.fc2 = nn.Linear(128, num_classes)  # Output one value for binary classification

    def forward(self, x):
        x = x.unsqueeze(1)  # Add channel dimension
        x = torch.relu(self.conv1(x))
        x = torch.relu(self.conv2(x))
        x = self.flatten(x)
        x = torch.relu(self.fc1(x))
        x = self.fc2(x)
        return x

# Load CNN Model
cnn = CNNModel(input_dim=X_train.shape[1])
cnn_model_path = r"E:\CloudAnomalyDetectionSystem\models\cnn_model.pth"
if not os.path.exists(cnn_model_path):
    raise FileNotFoundError(f"CNN model file not found at {cnn_model_path}")
cnn.load_state_dict(torch.load(cnn_model_path, weights_only=True))

# Function to calculate metrics for both normal and anomalous data
def calculate_metrics(y_true, y_pred):
    # Calculate confusion matrix
    tn, fp, fn, tp = confusion_matrix(y_true, y_pred).ravel()
    
    # Metrics for Normal Data (0)
    precision_normal = tn / (tn + fp) if (tn + fp) != 0 else 0
    recall_normal = tn / (tn + fn) if (tn + fn) != 0 else 0
    
    # Metrics for Anomalous Data (1)
    precision_anomalous = tp / (tp + fp) if (tp + fp) != 0 else 0
    recall_anomalous = tp / (tp + fn) if (tp + fn) != 0 else 0
    
    # Calculate F1-Score for both Normal and Anomalous
    f1_normal = f1_score(y_true, y_pred, pos_label=0)
    f1_anomalous = f1_score(y_true, y_pred, pos_label=1)
    
    # ROC AUC score
    roc_auc = roc_auc_score(y_true, y_pred)
    
    return {
        'precision_normal': precision_normal,
        'recall_normal': recall_normal,
        'f1_normal': f1_normal,
        'precision_anomalous': precision_anomalous,
        'recall_anomalous': recall_anomalous,
        'f1_anomalous': f1_anomalous,
        'roc_auc': roc_auc
    }

# Evaluate Autoencoder
autoencoder.eval()
with torch.no_grad():
    reconstructed = autoencoder(X_val_tensor)
    mse = torch.mean((reconstructed - X_val_tensor) ** 2, dim=1)
    threshold = mse.mean() + 3 * mse.std()
    ae_predictions = (mse > threshold).int()

# Calculate Autoencoder Metrics
ae_metrics = calculate_metrics(y_val, ae_predictions)

# Evaluate CNN
cnn.eval()
with torch.no_grad():
    outputs = cnn(X_val_tensor)
    outputs = torch.sigmoid(outputs)  # Apply sigmoid for binary classification
    cnn_predictions = (outputs > 0.5).int()  # Convert to binary predictions

# Calculate CNN Metrics
cnn_metrics = calculate_metrics(y_val, cnn_predictions)

# Save Evaluation Metrics to Word Document
doc = Document()
doc.add_heading('Evaluation Metrics for Anomaly Detection Models', 0)

# Autoencoder Metrics
doc.add_heading('Autoencoder Model', level=1)
doc.add_paragraph(f"Precision (Normal): {ae_metrics['precision_normal']:.4f}")
doc.add_paragraph(f"Recall (Normal): {ae_metrics['recall_normal']:.4f}")
doc.add_paragraph(f"F1-Score (Normal): {ae_metrics['f1_normal']:.4f}")
doc.add_paragraph(f"Precision (Anomalous): {ae_metrics['precision_anomalous']:.4f}")
doc.add_paragraph(f"Recall (Anomalous): {ae_metrics['recall_anomalous']:.4f}")
doc.add_paragraph(f"F1-Score (Anomalous): {ae_metrics['f1_anomalous']:.4f}")
doc.add_paragraph(f"ROC AUC: {ae_metrics['roc_auc']:.4f}")

# CNN Metrics
doc.add_heading('CNN Model', level=1)
doc.add_paragraph(f"Precision (Normal): {cnn_metrics['precision_normal']:.4f}")
doc.add_paragraph(f"Recall (Normal): {cnn_metrics['recall_normal']:.4f}")
doc.add_paragraph(f"F1-Score (Normal): {cnn_metrics['f1_normal']:.4f}")
doc.add_paragraph(f"Precision (Anomalous): {cnn_metrics['precision_anomalous']:.4f}")
doc.add_paragraph(f"Recall (Anomalous): {cnn_metrics['recall_anomalous']:.4f}")
doc.add_paragraph(f"F1-Score (Anomalous): {cnn_metrics['f1_anomalous']:.4f}")
doc.add_paragraph(f"ROC AUC: {cnn_metrics['roc_auc']:.4f}")

# Save Document
doc_path = os.path.join(results_dir, "evaluation(ae_cnn).docx")
doc.save(doc_path)
print(f"Evaluation metrics saved to {doc_path}")
